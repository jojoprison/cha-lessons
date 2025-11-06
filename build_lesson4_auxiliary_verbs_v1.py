import argparse
import json
# -*- coding: utf-8 -*-
# build_lesson4_auxiliary_verbs_v1.py
# Генерит DOCX: cha_lesson_4_auxiliary_verbs_v1.docx на основе cha_lesson_4_auxiliary_verbs_lite_v3.docx
# Требования:
# - Добавить RU строку после каждой EN строки в Explanation / Practice / Vocabulary Exercises / Exit check & Homework
# - В Vocabulary после RU добавить « — TH» перевод модальных/вспомогательных.
import os
import re
import time

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

# ---------- Цвета и стили ----------
GOLD = RGBColor(184, 134, 11)
BLACK = RGBColor(0, 0, 0)
DARK_RED = RGBColor(139, 0, 0)
DARK_GREEN = RGBColor(0, 100, 0)
PURPLE = RGBColor(102, 0, 153)

THAI_FONT_NAME = "Noto Sans Thai"

SRC_NAME = "cha_lesson_4_auxiliary_verbs_lite_v3.docx"
OUT_NAME = "cha_lesson_4_auxiliary_verbs_v1.docx"

# Тайский словарь для Vocabulary (ключ — нормализованный EN термин)
TH_VOCAB = {
    "can": "สามารถ",
    "could": "อาจจะ / สามารถ(อดีต)",
    "may": "อาจจะ",
    "might": "อาจจะ",
    "must": "ต้อง",
    "have to": "จำเป็นต้อง / ต้อง",
    "has to": "จำเป็นต้อง / ต้อง",
    "had to": "จำเป็นต้อง / ต้อง (อดีต)",
    "should": "ควร",
    "would": "จะ / มักจะ (สมมุติ)",
    "will": "จะ",
    "shall": "จะ (ทางการ)",
    "do": "ทำ (ตัวช่วยไวยากรณ์)",
    "does": "ทำ (ตัวช่วยไวยากรณ์)",
    "did": "ทำ (อดีต, ตัวช่วยไวยากรณ์)",
    "be": "เป็น/อยู่/คือ",
    "am": "เป็น/อยู่/คือ",
    "is": "เป็น/อยู่/คือ",
    "are": "เป็น/อยู่/คือ",
    "was": "เป็น/อยู่/คือ",
    "were": "เป็น/อยู่/คือ",
    "have": "มี / ได้ทำ (สมบูรณ์)",
    "has": "มี / ได้ทำ (สมบูรณ์)",
    "had": "มี / ได้ทำ (อดีต)",
    "be able to": "สามารถ",
    "need to": "จำเป็นต้อง",
    "ought to": "ควรจะ",
    "used to": "เคย",
    "dare": "กล้า",
    "had better": "ควรจะ...ดีกว่า",
}

# Тайский словарь для Word bank (School & Stationery)
WORD_BANK_TH = {
    "notebook": "สมุดโน้ต",
    "textbook": "ตำราเรียน",
    "workbook": "สมุดแบบฝึกหัด",
    "binder": "แฟ้มสันห่วง",
    "folder": "แฟ้ม",
    "loose-leaf paper": "กระดาษแยกแผ่น",
    "pen": "ปากกา",
    "pencil": "ดินสอ",
    "eraser": "ยางลบ",
    "sharpener": "กบเหลาดินสอ",
    "highlighter": "ปากกาเน้นข้อความ",
    "marker": "ปากกาเมจิก",
    "ruler": "ไม้บรรทัด",
    "protractor": "ไม้โปรแทรกเตอร์",
    "compass (geometry)": "วงเวียน",
    "glue stick": "กาวแท่ง",
    "scissors": "กรรไกร",
    "stapler": "ที่เย็บกระดาษ",
    "paper clips": "คลิปหนีบกระดาษ",
    "sticky notes": "กระดาษโพสต์อิท",
}

# Базовый RU-словарь для Vocabulary
RU_VOCAB = {
    "can": "может",
    "could": "мог(ла)/могли",
    "may": "может",
    "might": "возможно",
    "must": "должен",
    "have to": "должен/приходится",
    "has to": "должен/приходится",
    "had to": "должен был/пришлось",
    "should": "следует",
    "would": "бы",
    "will": "будет",
    "shall": "будет (офиц.)",
    "do": "делать (всп.)",
    "does": "делает (всп.)",
    "did": "сделал (всп.)",
    "be": "быть",
    "am": "есть",
    "is": "есть",
    "are": "есть",
    "was": "был",
    "were": "были",
    "have": "иметь",
    "has": "имеет",
    "had": "имел",
    "be able to": "может/в состоянии",
    "need to": "нужно/необходимо",
    "ought to": "следовало бы",
    "used to": "раньше делал/обычно",
    "dare": "осмеливаться",
    "had better": "лучше бы/следует",
}


# (автоперевод удалён; переводы берём только из источника)


def new_doc():
    doc = Document()
    for s in doc.sections:
        s.page_height = Cm(29.7)
        s.page_width = Cm(21.0)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        fp = s.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = fp.add_run("© Cha 2025 · Page ")
        run1.font.size = Pt(9)
        run1.font.color.rgb = BLACK
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run2 = fp.add_run()
        run2._r.append(fld)
    return doc


def clone_run(dst_p, src_run):
    r = dst_p.add_run(src_run.text)
    # базовые атрибуты
    try:
        r.font.bold = src_run.font.bold
        r.font.italic = src_run.font.italic
        r.font.underline = src_run.font.underline
        r.font.size = src_run.font.size
        r.font.all_caps = src_run.font.all_caps
        if src_run.font.color and src_run.font.color.rgb:
            r.font.color.rgb = src_run.font.color.rgb
    except Exception:
        pass
    return r


def clone_paragraph(dst_doc, src_p):
    p = dst_doc.add_paragraph()
    # копируем выравнивание, если нужно
    p.alignment = src_p.alignment
    for run in src_p.runs:
        clone_run(p, run)
    return p


# (удалены функции автоперевода; используем только заранее заданные переводы)


# (удалены неиспользуемые эвристики заголовков)


def is_vocab_item(text: str) -> bool:
    # Примитивная эвристика: строка типа "A. can — может"
    t = text.strip()
    if re.match(r"^[A-Za-z]\.\s+", t):
        return True
    # или строка с EN — RU уже
    if " — " in t and not any(ch.isdigit() for ch in t.split(" — ")[0][:3]):
        return True
    return False


def normalize_key(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().lower())


def norm_exact(s: str) -> str:
    # Нормализация ключей для словаря переводов (без нижнего регистра)
    s = (s or "").strip()
    # Унификация тире/дефисов и неразрывных дефисов
    s = s.replace("\u2011", "-")  # non-breaking hyphen
    s = s.replace("\u2013", "-")  # en dash
    s = s.replace("\u2014", "-")  # em dash
    s = s.replace("\u2212", "-")  # minus sign
    # Срезаем ведущую нумерацию вида 1)  1.  1.1  и т.п.
    s = re.sub(r"^\s*\d+(?:\.\d+)*[\)\.]?\s+", "", s)
    # Срезаем лидирующие маркеры списков (•, -, –, —) и пробелы
    s = re.sub(r"^[\u2022\-\u2013\u2014]\s+", "", s)
    # Упрощаем пробелы
    return re.sub(r"\s+", " ", s)


def strip_list_markers(s: str) -> str:
    return re.sub(r"^[\u2022\-\u2013\u2014]\s+", "", (s or "").strip())


def clean_vocab_en_term(s: str) -> str:
    """Очищает EN-термин в Word bank: убирает литерную нумерацию (a.), эмодзи, оставляет латиницу/пробелы/скобки/дефис."""
    s = (s or "").strip()
    # убрать a./b./c.
    s = re.sub(r"^[A-Za-z]\.[\s]+", "", s)
    # убрать эмодзи и прочие символы, кроме латиницы, пробелов, дефиса и ()
    s = re.sub(r"[^A-Za-z()\-\s]", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s.lower()


BLOCK_TITLES = {
    "✏️ Lesson 4 — Auxiliary Verbs — Vocabulary: School & Stationery",
    "👩‍🏫 Explanation",
    "🧠 Practice",
    "✍️ Examples:",
    "✏️ Vocabulary (School & Stationery)",
    "✏️ Vocabulary",
    "✏️ Vocabulary Exercises",
    "🧾 Exit check (5 quick items):",
    "🧾 Exit check & Homework",
}


def load_translations_json(path: str) -> dict:
    if not path or not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    # нормализуем ключи
    return {norm_exact(k): v for k, v in data.items()}


def load_translations_from_source(path: str) -> dict:
    """
    Парсит файл, где EN строка идёт отдельно, а ниже 2 строки в скобках — RU и TH.
    Игнорируем большие заголовки блоков и словарь Word bank.
    """
    if not path or not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        lines = [ln.rstrip("\n") for ln in f]
    tr = {}
    section = None
    i = 0
    while i < len(lines):
        L = lines[i].strip()
        if not L:
            i += 1
            continue
        # Определяем секцию по заголовкам
        low = L.lower()
        if L in BLOCK_TITLES:
            if "vocabulary" in low and "exercises" not in low:
                section = "vocab"
            elif "vocabulary exercises" in low:
                section = "vocab_ex"
            elif "practice" in low:
                section = "practice"
            elif "exit check" in low or "homework" in low:
                section = "exit"
            elif "explanation" in low:
                section = "expl"
            i += 1
            continue
        # Пропускаем шапки Word bank и сами элементы словаря — они не переводятся этим слоем
        if section == "vocab":
            i += 1
            continue
        # EN-строка — если следующая строка начинается на '(' — это RU, а следующая за ней — TH
        if not L.startswith("("):
            ru = th = None
            if i + 1 < len(lines) and lines[i + 1].strip().startswith("("):
                ru = lines[i + 1].strip()
                if ru.startswith("(") and ru.endswith(")"):
                    ru = ru[1:-1]
            if i + 2 < len(lines) and lines[i + 2].strip().startswith("("):
                th = lines[i + 2].strip()
                if th.startswith("(") and th.endswith(")"):
                    th = th[1:-1]
            if ru or th:
                base_key = norm_exact(L)
                val = {"ru": ru, "th": th}
                tr[base_key] = val
                # альтернативный ключ без маркеров списков
                alt = norm_exact(strip_list_markers(L))
                if alt != base_key:
                    tr[alt] = val
                i += 3
                continue
        i += 1
    return tr


def collect_highlight_tokens(src_p) -> list:
    """Собираем токены (верхний регистр) из EN-абзаца для зеркального подчеркивания в RU."""
    tokens = []
    for run in src_p.runs:
        t = run.text or ""
        # захватываем куски вида ALL CAPS (включая фразы с пробелами)
        for m in re.finditer(r"[A-Z][A-Z ]+[A-Z]", t):
            tok = m.group(0).strip()
            if tok not in tokens:
                tokens.append(tok)
        # если ран подчёркнут и без капса — попробуем захватить слово
        try:
            if run.font and run.font.underline and not any(
                    ch.isupper() for ch in t):
                # берем короткую метку до 15 символов
                w = t.strip()
                if 0 < len(w) <= 15 and w not in tokens:
                    tokens.append(w)
        except Exception:
            pass
    # сортируем по длине (длиннее вперёд), чтобы не разбивать составные токены
    tokens.sort(key=len, reverse=True)
    return tokens


def add_ru_mapped_line_with_highlights(doc, src_p, ru_text):
    """Рисуем RU строку из словаря, но зеркалим подчёркнутые/ALL CAPS токены из EN, если они встречаются в RU.
    Фон RU — тёмно-красный курсив; совпавшие токены — чёрный bold+underline (и без курсива).
    """
    p = doc.add_paragraph()
    # открывающая скобка
    r0 = p.add_run("(")
    r0.font.italic = True
    r0.font.color.rgb = DARK_RED

    hi = collect_highlight_tokens(src_p)
    s = ru_text or ""
    i = 0
    while i < len(s):
        hit_pos = None
        hit_tok = None
        # ищем ближайшее вхождение любого токена
        for tok in hi:
            j = s.find(tok, i)
            if j != -1 and (hit_pos is None or j < hit_pos):
                hit_pos = j
                hit_tok = tok
        if hit_pos is None:
            # хвост — обычный RU
            r = p.add_run(s[i:])
            r.font.italic = True
            r.font.color.rgb = DARK_RED
            break
        # прелюдия до токена
        if hit_pos > i:
            r = p.add_run(s[i:hit_pos])
            r.font.italic = True
            r.font.color.rgb = DARK_RED
        # сам токен — чёрный bold+underline
        r2 = p.add_run(s[hit_pos:hit_pos + len(hit_tok)])
        r2.font.color.rgb = BLACK
        r2.font.bold = True
        r2.font.underline = True
        r2.font.italic = False
        i = hit_pos + len(hit_tok)

    # закрывающая скобка
    rz = p.add_run(")")
    rz.font.italic = True
    rz.font.color.rgb = DARK_RED


def add_th_mapped_line_with_highlights(doc, src_p, th_text):
    """TH строка с зеркалированием подчёркнутых/ALL CAPS токенов из EN.
    База — зелёный курсив; совпавшие токены — чёрный bold+underline (без курсива).
    """
    p = doc.add_paragraph()
    r0 = p.add_run("(")
    r0.font.italic = True
    r0.font.color.rgb = DARK_GREEN

    hi = collect_highlight_tokens(src_p)
    s = th_text or ""
    i = 0
    while i < len(s):
        hit_pos = None
        hit_tok = None
        for tok in hi:
            j = s.find(tok, i)
            if j != -1 and (hit_pos is None or j < hit_pos):
                hit_pos = j
                hit_tok = tok
        if hit_pos is None:
            r = p.add_run(s[i:])
            r.font.italic = True
            r.font.color.rgb = DARK_GREEN
            r.font.name = THAI_FONT_NAME
            break
        if hit_pos > i:
            r = p.add_run(s[i:hit_pos])
            r.font.italic = True
            r.font.color.rgb = DARK_GREEN
            r.font.name = THAI_FONT_NAME
        r2 = p.add_run(s[hit_pos:hit_pos + len(hit_tok)])
        r2.font.color.rgb = BLACK
        r2.font.bold = True
        r2.font.underline = True
        r2.font.italic = False
        r2.font.name = THAI_FONT_NAME
        i = hit_pos + len(hit_tok)

    rz = p.add_run(")")
    rz.font.italic = True
    rz.font.color.rgb = DARK_GREEN


def append_th_to_vocab_line(dst_p):
    # Разбираем текущую строку, пытаемся получить EN термин
    full = dst_p.text
    parts = full.split(" — ")
    # EN часть до первого тире или вся строка
    en_part = parts[0] if parts else full
    en_term = re.sub(r"^[A-Za-z]\.\s+", "", en_part).strip()

    # Подбор переводов
    # Сначала пробуем word bank (stationery)
    cleaned = clean_vocab_en_term(en_term)
    th = WORD_BANK_TH.get(cleaned)
    if not th:
        th = TH_VOCAB.get(normalize_key(en_term))
    if not th:
        for k in list(TH_VOCAB.keys()):
            if normalize_key(k) == normalize_key(en_term):
                th = TH_VOCAB[k]
                break
    ru = RU_VOCAB.get(normalize_key(en_term))
    if not ru:
        for k in list(RU_VOCAB.keys()):
            if normalize_key(k) == normalize_key(en_term):
                ru = RU_VOCAB[k]
                break

    # Если RU уже есть в строке — просто добавляем TH
    ru_added = False
    th_added = False
    if " — " in full:
        if th:
            rr = dst_p.add_run(" — ")
            rr.font.italic = True
            rr.font.color.rgb = DARK_GREEN
            tr = dst_p.add_run(th)
            tr.font.italic = True
            tr.font.color.rgb = DARK_GREEN
            tr.font.name = THAI_FONT_NAME
            th_added = True
        else:
            # лог пропуска TH для словарной строки
            try:
                print(f"[lesson4][miss][Vocab TH] {cleaned or en_term}")
            except Exception:
                pass
        return ru_added, th_added

    # Если RU не было — добавляем RU и TH
    if ru:
        rr_sep = dst_p.add_run(" — ")
        rr_sep.font.italic = True
        rr_sep.font.color.rgb = DARK_RED
        rr_run = dst_p.add_run(ru)
        rr_run.font.italic = True
        rr_run.font.color.rgb = DARK_RED
        ru_added = True
    if th:
        th_sep = dst_p.add_run(" — ")
        th_sep.font.italic = True
        th_sep.font.color.rgb = DARK_GREEN
        tr = dst_p.add_run(th)
        tr.font.italic = True
        tr.font.color.rgb = DARK_GREEN
        tr.font.name = THAI_FONT_NAME
        th_added = True
    else:
        try:
            print(f"[lesson4][miss][Vocab TH] {cleaned or en_term}")
        except Exception:
            pass
    return ru_added, th_added


def build():
    parser = argparse.ArgumentParser()
    parser.add_argument("--with-ru", dest="with_ru", action="store_true",
                        default=True)
    parser.add_argument("--no-ru", dest="with_ru", action="store_false")
    parser.add_argument("--with-th", dest="with_th", action="store_true",
                        default=True)
    parser.add_argument("--no-th", dest="with_th", action="store_false")
    # Отдельные флаги для Vocabulary
    parser.add_argument("--vocab-th", dest="vocab_th", action="store_true",
                        default=True)
    parser.add_argument("--no-vocab-th", dest="vocab_th", action="store_false")
    parser.add_argument("--vocab-ru", dest="vocab_ru", action="store_true",
                        default=False)
    parser.add_argument("--no-vocab-ru", dest="vocab_ru", action="store_false")
    parser.add_argument("--translations", type=str,
                        default="lesson4_translations.json")
    parser.add_argument("--translations-source", type=str,
                        default="lesson4_translations_source.txt")
    # (fallback авто-перевода удалён)
    args = parser.parse_args()
    start_ts = time.time()
    print("[lesson4] Start generation")
    src_path = os.path.join(os.getcwd(), SRC_NAME)
    if not os.path.exists(src_path):
        raise FileNotFoundError(f"Source DOCX not found: {src_path}")
    print(f"[lesson4] Source: {SRC_NAME}")
    src = Document(src_path)

    out = new_doc()
    print("[lesson4] New document initialized")
    # Грузим переводы (из source .txt приоритетнее, затем .json)
    tr_map = {}
    if args.translations_source and os.path.exists(args.translations_source):
        tr_map = load_translations_from_source(args.translations_source)
        print(
            f"[lesson4] Translations loaded from: {args.translations_source} ({len(tr_map)} entries)")
    if not tr_map and args.translations and os.path.exists(args.translations):
        tr_map = load_translations_json(args.translations)
        print(
            f"[lesson4] Translations loaded from: {args.translations} ({len(tr_map)} entries)")

    # Простая машина состояний по секциям
    section = None
    ru_lines = 0
    vocab_th_added = 0
    vocab_ru_added = 0

    total = len(src.paragraphs)
    print(f"[lesson4] Paragraphs: {total}")

    for idx, p in enumerate(src.paragraphs, 1):
        text = p.text or ""
        # Клонируем исходную строку как есть
        new_p = clone_paragraph(out, p)

        # Определяем смену секции по ключевым словам
        t = text.strip().lower()
        if "vocabulary" in t and len(t) < 64:
            if section != "vocab":
                print("[lesson4] --> Section: Vocabulary")
            section = "vocab"
        elif "vocabulary exercises" in t:
            if section != "vocab_ex":
                print("[lesson4] --> Section: Vocabulary Exercises")
            section = "vocab_ex"
        elif "practice" in t:
            if section != "practice":
                print("[lesson4] --> Section: Practice")
            section = "practice"
        elif "exit check" in t or "homework" in t:
            if section != "exit":
                print("[lesson4] --> Section: Exit check & Homework")
            section = "exit"
        elif "explanation" in t or "examples" in t:
            if section != "expl":
                print("[lesson4] --> Section: Explanation/Examples")
            section = "expl"

        # Если это пункт словаря — добавим TH
        if section == "vocab" and is_vocab_item(text):
            before = new_p.text
            # функция вернёт, добавляли ли RU/TH
            ru_added = False
            th_added = False
            if args.vocab_th:
                _, th_added = append_th_to_vocab_line(new_p)
            # при необходимости можно добавить RU к словарю, пока по умолчанию выключено
            if args.vocab_ru and not ru_added:
                # RU обычно уже присутствует в лексике, поэтому здесь пропускаем
                pass
            if ru_added:
                vocab_ru_added += 1
            if th_added:
                vocab_th_added += 1
            continue  # для словаря RU-строку отдельную не вставляем (она уже на линии)

        # Для всех остальных контентных EN строк — добавляем RU строку
        stripped = text.strip()
        if not stripped:
            continue
        # Для заголовков разделов не добавляем (не переводим названия блоков)
        if stripped in BLOCK_TITLES:
            continue

        # Ищем переводы в словаре (приоритетнее авто-перевода). Применяем по флагам.
        key = norm_exact(text)
        has_any = False
        if args.with_ru:
            ru_txt = tr_map.get(key, {}).get("ru")
            if ru_txt:
                add_ru_mapped_line_with_highlights(out, p, ru_txt)
                ru_lines += 1
                has_any = True
            else:
                # короткий лог по пропускам RU
                print(f"[lesson4][miss][RU] {key[:80]}")
        if args.with_th:
            th_txt = tr_map.get(key, {}).get("th")
            if th_txt:
                add_th_mapped_line_with_highlights(out, p, th_txt)
                has_any = True
            else:
                print(f"[lesson4][miss][TH] {key[:80]}")
        # Перевод добавляется только если есть в словаре; авто-перевод отключён

        # прогресс каждые 20 параграфов
        if idx % 20 == 0:
            print(f"[lesson4] Progress: {idx}/{total} paragraphs processed")

    out.save(OUT_NAME)
    dur = time.time() - start_ts
    print("[lesson4] Saved:", OUT_NAME)
    print(
        f"[lesson4] Summary: RU lines added={ru_lines}, vocab TH added={vocab_th_added}, vocab RU added={vocab_ru_added}")
    print(f"[lesson4] Done in {dur:.1f}s")


if __name__ == "__main__":
    build()
