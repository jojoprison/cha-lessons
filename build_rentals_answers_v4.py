# -*- coding: utf-8 -*-
# build_rentals_answers_v4.py
# Генерит DOCX: cha_test_12tenses_rentals_with_answers_v4.docx
# Формат — наш «Cha» (EN — gold/bold; RU — dark red italic; TH — dark green italic; ответы — purple)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

# ---------- Цвета и стили ----------
GOLD = RGBColor(184, 134, 11)  # EN (жёлтый/золото)
BLACK = RGBColor(0, 0, 0)
DARK_RED = RGBColor(139, 0, 0)  # RU
DARK_GREEN = RGBColor(0, 100, 0)  # TH
PURPLE = RGBColor(102, 0, 153)  # Ответы/объяснения (EN only)

THAI_FONT_NAME = "Noto Sans Thai"  # если установишь локально — шрифт подхватится


def new_doc():
    doc = Document()
    for s in doc.sections:
        s.page_height = Cm(29.7)
        s.page_width = Cm(21.0)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        # Footer: © Cha 2025 + номер страницы
        fp = s.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = fp.add_run("© Cha 2025 · Page ")
        run1.font.size = Pt(9)
        run1.font.color.rgb = BLACK
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run2 = fp.add_run()
        run2._r.append(fld)
    return doc


def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLACK


def add_section_title(doc, emoji, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{emoji} {text}")
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BLACK


def add_blank(doc, lines=1):
    for _ in range(lines):
        doc.add_paragraph("")


def run_gold(p, txt, size=12, bold=True):
    r = p.add_run(txt)
    r.font.color.rgb = GOLD
    r.font.bold = bold
    r.font.size = Pt(size)
    return r


def run_black(p, txt, size=12, bold=True, underline=False):
    r = p.add_run(txt)
    r.font.color.rgb = BLACK
    r.font.bold = bold
    r.font.underline = underline
    r.font.size = Pt(size)
    return r


def line_ru(doc, txt, size=11):
    p = doc.add_paragraph()
    r = p.add_run(f"({txt})")
    r.font.italic = True
    r.font.color.rgb = DARK_RED
    r.font.size = Pt(size)


def line_th(doc, txt, size=11):
    p = doc.add_paragraph()
    r = p.add_run(f"({txt})")
    r.font.italic = True
    r.font.color.rgb = DARK_GREEN
    r.font.size = Pt(size)
    r.font.name = THAI_FONT_NAME


def prompt_line(p, text):
    """
    Строка-подсказка: основной текст (EN) — gold,
    куски в [[...]] — подчёркнутые чёрные (что нужно выбрать/вставить).
    """
    i = 0
    while i < len(text):
        s = text.find("[[", i)
        if s == -1:
            run_gold(p, text[i:])
            break
        if s > i:
            run_gold(p, text[i:s])
        e = text.find("]]", s + 2)
        if e == -1:
            e = len(text)
        chunk = text[s + 2:e]
        run_black(p, chunk, underline=True, bold=True)
        i = e + 2


def examples_block(doc, items, idx_prefix):
    p = doc.add_paragraph()
    run_black(p, "✍️ Examples:", bold=True)
    for j, ex in enumerate(items, 1):
        p2 = doc.add_paragraph()
        run_black(p2, f"{idx_prefix}.{j} ", bold=True)
        # сами примеры — подчёркнутые чёрные, чтобы визуально отделялись
        r = run_black(p2, ex, bold=False, underline=True)


def add_word_bank_item(doc, letter, en, ru=None, th=None, emoji=None):
    p = doc.add_paragraph()
    # буква-индекс — чёрная
    idx = p.add_run(f"{letter}. ")
    idx.font.bold = True
    idx.font.color.rgb = BLACK
    idx.font.size = Pt(12)
    # эмодзи (если есть)
    if emoji:
        ee = p.add_run(f"{emoji} ")
        ee.font.bold = True
        ee.font.size = Pt(12)
    # английское слово — жёлтое, bold
    en_run = p.add_run(en)
    en_run.font.bold = True
    en_run.font.color.rgb = GOLD
    en_run.font.size = Pt(12)
    # RU — тёмно-красный курсив
    if ru:
        p.add_run(" — ")
        rr = p.add_run(ru)
        rr.font.italic = True
        rr.font.color.rgb = DARK_RED
    # TH — тёмно-зелёный курсив
    if th:
        p.add_run(" — ")
        tt = p.add_run(th)
        tt.font.italic = True
        tt.font.color.rgb = DARK_GREEN
        tt.font.name = THAI_FONT_NAME


def add_exercise(doc, idx, en, ru=None, th=None):
    p = doc.add_paragraph()
    run_black(p, f"{idx}) ", bold=True)  # номер — чёрный bold
    prompt_line(p, en)  # основная строка (жёлтая) + чёрные подчёркнутые вставки
    if ru:
        line_ru(doc, ru)
    if th:
        line_th(doc, th)


def add_answer_block(doc, answer_en, explanation_en, explanation_th=None):
    """
    Ответы: фиолетовый (EN-only) Answer: <...> — <explanation>
    Переводы RU/TH идут отдельными строками выше (как и в упражнении).
    """
    p = doc.add_paragraph()
    a = p.add_run("Answer: ")
    a.font.bold = True
    a.font.color.rgb = PURPLE
    b = p.add_run(f"{answer_en} — {explanation_en}")
    b.font.color.rgb = PURPLE
    if explanation_th:
        p2 = doc.add_paragraph()
        t = p2.add_run(explanation_th)
        t.font.color.rgb = PURPLE
        t.font.italic = True


# ---------- Контент (как согласовали) ----------
content = {
    "title": "🏠 Control Test — 12 Tenses (Rentals & Housing) — v4",
    "theme_emoji": "🧰",
    "vocab_title": "Vocabulary (Rentals & Housing)",
    "explanation": [
        {"title": "1) ✅ Present Simple — habits/routines. Form: V1 / V1+s.",
         "ru": "Настоящее простое — привычки/распорядок. Формула: V1 / V1+s.",
         "th": "ปัจจุบันธรรมดา — พฤติกรรม/กิจวัตร รูป: V1 / V1+s",
         "examples": ["Tenants pay rent monthly.", "The landlord checks IDs.",
                      "We sign contracts here."], "i": "1"},
        {"title": "2) ⏳ Past Simple — finished past event. Form: V2 / V-ed.",
         "ru": "Прошедшее простое — завершённое действие в прошлом. V2 / V-ed.",
         "th": "อดีตกาลธรรมดา — เหตุการณ์จบในอดีต รูป: V2 / V-ed",
         "examples": ["She moved out last week.",
                      "They paid the deposit yesterday.",
                      "I called the agent."], "i": "2"},
        {"title": "3) 🔮 Future Simple — decisions/promises. Form: will + V1.",
         "ru": "Будущее простое — решения/обещания. will + V1.",
         "th": "อนาคตกาลธรรมดา — การตัดสินใจ/สัญญา will + V1",
         "examples": ["I will renew the lease.", "We will fix the sink.",
                      "The owner will respond soon."], "i": "3"},
        {
            "title": "4) 🔄 Present Continuous — happening now/temporary. Form: am/is/are + V-ing.",
            "ru": "Настоящее продолженное — сейчас/временное. am/is/are + V-ing.",
            "th": "ปัจจุบันต่อเนื่อง — ขณะนี้/ชั่วคราว am/is/are + V-ing",
            "examples": ["The plumber is repairing the pipe.",
                         "I am showing the apartment.",
                         "They are painting the walls."], "i": "4"},
        {
            "title": "5) ⏪ Past Continuous — in progress at a past time. Form: was/were + V-ing.",
            "ru": "Прошедшее продолженное — процесс в прошлом. was/were + V-ing.",
            "th": "อดีตต่อเนื่อง — กำลังกระทำในอดีต was/were + V-ing",
            "examples": ["We were signing papers at 3 pm.",
                         "She was cleaning the flat.",
                         "They were moving furniture."], "i": "5"},
        {
            "title": "6) 🛰️ Future Continuous — in progress at a future time. Form: will be + V-ing.",
            "ru": "Будущее продолженное — процесс в будущем. will be + V-ing.",
            "th": "อนาคตต่อเนื่อง — กำลังกระทำในอนาคต will be + V-ing",
            "examples": ["Tomorrow I will be meeting the landlord.",
                         "They will be inspecting units.",
                         "We will be moving in at noon."], "i": "6"},
        {"title": "7) 🏁 Present Perfect — result now. Form: have/has + V3.",
         "ru": "Настоящее совершённое — результат к настоящему. have/has + V3.",
         "th": "ปัจจุบันสมบูรณ์ — ผลลัพธ์ตอนนี้ have/has + V3",
         "examples": ["I have paid the deposit.", "They have signed the lease.",
                      "He has fixed the door."], "i": "7"},
        {"title": "8) 🕰️ Past Perfect — earlier past. Form: had + V3.",
         "ru": "Предпрошедшее — действие раньше другого прошлого. had + V3.",
         "th": "อดีตก่อนอดีต — เกิดก่อนอีกเหตุการณ์ในอดีต had + V3",
         "examples": ["We had moved in before winter.",
                      "She had reported the issue.",
                      "They had checked IDs already."], "i": "8"},
        {
            "title": "9) 🚀 Future Perfect — completed by a future point. Form: will have + V3.",
            "ru": "Будущее совершённое — будет завершено к моменту. will have + V3.",
            "th": "อนาคตสมบูรณ์ — เสร็จสิ้นก่อนจุดเวลาอนาคต will have + V3",
            "examples": ["By Friday, we will have painted the room.",
                         "I will have paid all bills.",
                         "They will have renewed the lease."], "i": "9"},
        {
            "title": "10) 🌱 Present Perfect Continuous — duration to now. Form: have/has been + V-ing.",
            "ru": "Наст. сов. продолж. — длительность до настоящего. have/has been + V-ing.",
            "th": "ปัจจุบันสมบูรณ์ต่อเนื่อง — ระยะเวลาถึงปัจจุบัน have/has been + V-ing",
            "examples": ["We have been looking for a flat for weeks.",
                         "She has been waiting at the office.",
                         "They have been renovating the kitchen."], "i": "10"},
        {
            "title": "11) 🧭 Past Perfect Continuous — duration before past point. Form: had been + V-ing.",
            "ru": "Предпрош. продолж. — длительность до прошлого момента. had been + V-ing.",
            "th": "อดีตก่อนอดีตต่อเนื่อง — ระยะเวลาก่อนจุดในอดีต had been + V-ing",
            "examples": ["I had been negotiating for months.",
                         "They had been living there since 2019.",
                         "He had been fixing leaks all day."], "i": "11"},
        {
            "title": "12) 🔭 Future Perfect Continuous — duration until future point. Form: will have been + V-ing.",
            "ru": "Буд. сов. продолж. — длит. к будущему моменту. will have been + V-ing.",
            "th": "อนาคตสมบูรณ์ต่อเนื่อง — ระยะเวลาถึงจุดเวลาอนาคต will have been + V-ing",
            "examples": ["By June, we will have been renting for a year.",
                         "She will have been managing the unit for months.",
                         "They will have been renovating for weeks."],
            "i": "12"},
    ],
    # Practice — 50 заданий. Партия 1/3 (1–18) из ваших листов
    "practice": [
        {"en": "The tenant usually [[(pays)]] rent on the first of each month.",
         "th": "ผู้เช่ามักจ่ายค่าเช่าในวันแรกของทุกเดือน",
         "answer_en": "pays", "explain_en": "Present Simple (habit/routine).",
         "explain_th": "คำอธิบาย: ใช้ Present Simple (นิสัย/กิจวัตร)."},
        {
            "en": "Right now, the landlord [[(is inspecting)]] the apartment for any damage.",
            "th": "ตอนนี้เจ้าของบ้านกำลังตรวจอพาร์ตเมนต์หาความเสียหาย",
            "answer_en": "is inspecting",
            "explain_en": "Present Continuous (action in progress now).",
            "explain_th": "คำอธิบาย: Present Continuous ใช้กับการกระทำที่กำลังเกิดขึ้นขณะนี้."},
        {
            "en": "Yesterday at 6 p.m., we [[(were signing)]] the lease at the agency.",
            "th": "เมื่อวานเวลา 6 โมงเย็น เรากำลังเซ็นสัญญาที่เอเจนซี่",
            "answer_en": "were signing",
            "explain_en": "Past Continuous (in progress at a past time).",
            "explain_th": "คำอธิบาย: Past Continuous แสดงเหตุการณ์ที่กำลังดำเนินอยู่ในเวลาหนึ่งในอดีต."},
        {
            "en": "By tomorrow morning, the agent [[(will have prepared)]] the renewal contract.",
            "th": "ภายในเช้าวันพรุ่งนี้ ตัวแทนจะได้เตรียมสัญญาต่ออายุเสร็จแล้ว",
            "answer_en": "will have prepared",
            "explain_en": "Future Perfect (completed before a future point).",
            "explain_th": "คำอธิบาย: Future Perfect เสร็จก่อนจุดเวลาในอนาคต."},
        {"en": "They [[(have already moved in)]] and unpacked most boxes.",
         "th": "พวกเขาย้ายเข้าแล้วและแกะกล่องส่วนใหญ่แล้ว",
         "answer_en": "have already moved in",
         "explain_en": "Present Perfect (result now).",
         "explain_th": "คำอธิบาย: Present Perfect ใช้กับผลลัพธ์ที่มีผลถึงปัจจุบัน (already)."},
        {
            "en": "Before we got the keys, the cleaners [[(had finished)]] the deep clean.",
            "th": "ก่อนที่พวกเราจะได้รับกุญแจ แม่บ้านได้ทำความสะอาดใหญ่เสร็จแล้ว",
            "answer_en": "had finished",
            "explain_en": "Past Perfect (earlier past action).",
            "explain_th": "คำอธิบาย: Past Perfect เสร็จก่อนเหตุการณ์อีกเหตุการณ์หนึ่งในอดีต."},
        {
            "en": "At 8 p.m. tonight, we [[(will be discussing)]] the pet policy with the landlord.",
            "th": "คืนนี้เวลา 2 ทุ่ม เราจะกำลังคุยเรื่องนโยบายสัตว์เลี้ยงกับเจ้าของบ้าน",
            "answer_en": "will be discussing",
            "explain_en": "Future Continuous (in progress at a future time).",
            "explain_th": "คำอธิบาย: Future Continuous กำลังกระทำ ณ เวลาในอนาคต."},
        {
            "en": "For three months, the tenant [[(has been waiting)]] for a plumbing repair.",
            "th": "เป็นเวลา 3 เดือนแล้ว ผู้เช่ากำลังรอการซ่อมท่อประปา",
            "answer_en": "has been waiting",
            "explain_en": "Present Perfect Continuous (duration to now).",
            "explain_th": "คำอธิบาย: Present Perfect Continuous ใช้กับความต่อเนื่องถึงปัจจุบัน."},
        {
            "en": "By last winter, we [[(had been paying)]] higher utilities for a year.",
            "th": "ถึงฤดูหนาวปีที่แล้ว เราได้จ่ายค่าสาธารณูปโภคสูงขึ้นมาแล้วหนึ่งปี",
            "answer_en": "had been paying",
            "explain_en": "Past Perfect Continuous (duration to a past point).",
            "explain_th": "คำอธิบาย: PPC แสดงความต่อเนื่องก่อนจุดเวลาในอดีต."},
        {
            "en": "Next July, they [[(will have been living)]] here for two years.",
            "th": "เดือนกรกฎาคมหน้า พวกเขาจะอาศัยอยู่ที่นี่ครบสองปี",
            "answer_en": "will have been living",
            "explain_en": "Future Perfect Continuous (duration to a future point).",
            "explain_th": "คำอธิบาย: FPC ใช้ระยะเวลาถึงจุดเวลาในอนาคต."},
        {
            "en": "When the inspector arrived, the contractor [[(had already fixed)]] the smoke detector.",
            "th": "เมื่อผู้ตรวจมาถึง ผู้รับเหมาได้ซ่อมเครื่องตรวจจับควันเสร็จแล้ว",
            "answer_en": "had already fixed",
            "explain_en": "Past Perfect (completed before past arrival).",
            "explain_th": "คำอธิบาย: Past Perfect เสร็จก่อนเหตุการณ์ในอดีต (already)."},
        {
            "en": "I can't find the mailbox key — I [[(have misplaced)]] it again.",
            "th": "ฉันหากุญแจตู้จดหมายไม่เจอ — ฉันวางผิดที่อีกแล้ว",
            "answer_en": "have misplaced",
            "explain_en": "Present Perfect (recent result).",
            "explain_th": "คำอธิบาย: Present Perfect ใช้กับผลลัพธ์ล่าสุด."},
        {
            "en": "At noon tomorrow, maintenance [[(will be checking)]] the heating system.",
            "th": "พรุ่งนี้เที่ยง ทีมซ่อมบำรุงจะกำลังตรวจระบบทำความร้อน",
            "answer_en": "will be checking", "explain_en": "Future Continuous.",
            "explain_th": "คำอธิบาย: Future Continuous."},
        {
            "en": "Before we move out, we [[(will have completed)]] the condition report.",
            "th": "ก่อนที่เราจะย้ายออก เราจะทำรายงานสภาพเสร็จแล้ว",
            "answer_en": "will have completed", "explain_en": "Future Perfect.",
            "explain_th": "คำอธิบาย: Future Perfect — เสร็จก่อนย้ายออก."},
        {
            "en": "By the time the realtor called, I [[(had been reviewing)]] listings for hours.",
            "th": "ตอนที่นายหน้าโทร ฉันทบทวนลิสต์มาเป็นชั่วโมงแล้ว",
            "answer_en": "had been reviewing",
            "explain_en": "Past Perfect Continuous.",
            "explain_th": "คำอธิบาย: PPC — ระยะเวลาต่อเนื่องก่อนเหตุการณ์ในอดีต."},
        {"en": "Since January, we [[(have been negotiating)]] a lower rent.",
         "th": "ตั้งแต่มกราคม เราได้ต่อรองค่าเช่าที่ถูกลงอย่างต่อเนื่อง",
         "answer_en": "have been negotiating",
         "explain_en": "Present Perfect Continuous.",
         "explain_th": "คำอธิบาย: PPC — since/for แสดงความต่อเนื่องถึงตอนนี้."},
        {
            "en": "Look — the handyman [[(is repairing)]] the leaking sink right now.",
            "th": "ดูสิ — ช่างกำลังซ่อมอ่างล้างมือที่รั่วอยู่ตอนนี้",
            "answer_en": "is repairing",
            "explain_en": "Present Continuous (right now).",
            "explain_th": "คำอธิบาย: Present Continuous — right now."},
        {
            "en": "Every spring the landlord [[(raises)]] the rent unless we negotiate.",
            "th": "ทุกฤดูใบไม้ผลิ เจ้าของบ้านขึ้นค่าเช่า เว้นแต่เราจะต่อรอง",
            "answer_en": "raises", "explain_en": "Present Simple (habit).",
            "explain_th": "คำอธิบาย: Present Simple — กิจวัตร/นิสัย."},
        {
            "en": "By last week's inspection, they [[(had already painted)]] the living room.",
            "th": "ถึงวันตรวจเมื่อสัปดาห์ที่แล้ว พวกเขาทาสีห้องนั่งเล่นเสร็จแล้ว",
            "answer_en": "had already painted",
            "explain_en": "Past Perfect (completed before a past point).",
            "explain_th": "คำอธิบาย: Past Perfect — เสร็จก่อนเหตุการณ์ในอดีต."},
        {"en": "Next time we talk, I [[(will have signed)]] the renewal.",
         "th": "ครั้งหน้าที่เราคุยกัน ฉันจะเซ็นต่อสัญญาเสร็จแล้ว",
         "answer_en": "will have signed",
         "explain_en": "Future Perfect (completed before the next time).",
         "explain_th": "คำอธิบาย: Future Perfect — เสร็จก่อนเวลาที่กำหนด."},
        {
            "en": "By 8 p.m., we [[(will have finished)]] the move-out cleaning and [[(will be handing)]] over the keys.",
            "th": "ภายในสองทุ่ม เราจะทำความสะอาดย้ายออกเสร็จ และจะกำลังส่งมอบกุญแจ",
            "answer_en": "will have finished + will be handing",
            "explain_en": "Future Perfect + Future Continuous.",
            "explain_th": "คำอธิบาย: เสร็จหนึ่งอย่าง (FPerf) และอีกอย่างกำลังทำอยู่ ณ เวลานั้น (FCont)."},
        {
            "en": "They [[(had been waiting)]] for approval while the manager [[(was reviewing)]] their references.",
            "th": "พวกเขารอการอนุมัติอยู่ ขณะผู้จัดการกำลังตรวจสอบเอกสารอ้างอิง",
            "answer_en": "had been waiting + was reviewing",
            "explain_en": "Past Perfect Continuous + Past Continuous.",
            "explain_th": "คำอธิบาย: PPC (รออย่างต่อเนื่อง) + PC (อีกเหตุการณ์กำลังเกิดขึ้น)."},
        {
            "en": "We [[(have been saving)]] for the deposit because the landlord [[(has increased)]] it this year.",
            "th": "เราเก็บเงินสำหรับมัดจำเพราะเจ้าของบ้านได้เพิ่มจำนวนในปีนี้",
            "answer_en": "have been saving + has increased",
            "explain_en": "Present Perfect Continuous + Present Perfect.",
            "explain_th": "คำอธิบาย: PPC (กระทำต่อเนื่อง) + PresPerf (การเปลี่ยนแปลงในปีนี้)."},
        {
            "en": "While the agent [[(was drafting)]] the addendum, we [[(were discussing)]] the notice period.",
            "th": "ขณะเอเจนต์กำลังร่างภาคผนวก เรากำลังหารือเรื่องระยะเวลาแจ้งออก",
            "answer_en": "was drafting + were discussing",
            "explain_en": "two Past Continuous actions.",
            "explain_th": "คำอธิบาย: Past Continuous สองเหตุการณ์เกิดควบคู่."},
        {
            "en": "By the time we get there, the tenants [[(will have been moving)]] furniture for three hours.",
            "th": "ตอนที่เราไปถึง ผู้เช่าจะย้ายของมาแล้วเป็นเวลา 3 ชั่วโมง",
            "answer_en": "will have been moving",
            "explain_en": "Future Perfect Continuous (duration to a future point).",
            "explain_th": "คำอธิบาย: FPC — ระยะเวลาถึงจุดเวลาในอนาคต."},
        {
            "en": "After the leak, the building [[(was being repaired)]] for weeks while we [[(were staying)]] in a hotel.",
            "th": "หลังมีน้ำรั่ว อาคารถูกซ่อมอยู่นานหลายสัปดาห์ ขณะที่เราพักโรงแรม",
            "answer_en": "was being repaired + were staying",
            "explain_en": "Past Continuous (passive) + Past Continuous (active).",
            "explain_th": "คำอธิบาย: Past Continuous รูปถูกกระทำ + รูปปกติพร้อมกัน."},
        {
            "en": "We [[(will be meeting)]] the property manager, and by then he [[(will have reviewed)]] our application.",
            "th": "เราจะพบผู้จัดการทรัพย์สิน และตอนนั้นเขาจะได้ตรวจใบสมัครของเราเสร็จแล้ว",
            "answer_en": "will be meeting + will have reviewed",
            "explain_en": "Future Continuous + Future Perfect.",
            "explain_th": "คำอธิบาย: กำลังกระทำ ณ เวลาอนาคต + เสร็จก่อนเวลานั้น."},
        {
            "en": "For two years they [[(have been renting)]] a furnished flat, but last month they [[(moved)]] to an unfurnished one.",
            "th": "ตลอดสองปี พวกเขาเช่าห้องที่มีเฟอร์นิเจอร์ แต่เดือนที่แล้วได้ย้ายไปห้องไม่มีเฟอร์นิเจอร์",
            "answer_en": "have been renting + moved",
            "explain_en": "Present Perfect Continuous + Past Simple.",
            "explain_th": "คำอธิบาย: PPC (ระยะเวลาต่อเนื่อง) + Past Simple (เหตุการณ์จุดเดียวในอดีต)."},
        {
            "en": "Before we viewed the place, the agent [[(had uploaded)]] new photos and [[(had updated)]] the listing.",
            "th": "ก่อนที่เราจะไปดูที่ เอเจนต์ได้อัปโหลดรูปใหม่และอัปเดตประกาศแล้ว",
            "answer_en": "had uploaded + had updated",
            "explain_en": "two Past Perfect actions.",
            "explain_th": "คำอธิบาย: Past Perfect สองเหตุการณ์ที่เกิดก่อน."},
        {
            "en": "Next week at 5, the handyman [[(will be installing)]] a new lock, and by 6 he [[(will have finished)]] the job.",
            "th": "สัปดาห์หน้าตอน 5 โมง ช่างจะกำลังติดตั้งกุญแจใหม่ และตอน 6 โมงเขาจะทำงานเสร็จแล้ว",
            "answer_en": "will be installing + will have finished",
            "explain_en": "Future Continuous + Future Perfect.",
            "explain_th": "คำอธิบาย: ทำอยู่ ณ เวลาในอนาคต และเสร็จภายในเวลาถัดมา."},
        {
            "en": "I [[(have already sent)]] the notice to vacate, so we [[(are packing)]] now.",
            "th": "ฉันได้ส่งหนังสือแจ้งย้ายออกแล้ว ดังนั้นตอนนี้เรากำลังเก็บของ",
            "answer_en": "have already sent + are packing",
            "explain_en": "Present Perfect + Present Continuous.",
            "explain_th": "คำอธิบาย: ผลลัพธ์ที่เกิดขึ้นแล้ว + การกระทำกำลังดำเนินอยู่ตอนนี้."},
        {
            "en": "By last Friday, we [[(had been waiting)]] for the deposit refund for two weeks.",
            "th": "ถึงวันศุกร์ที่แล้ว เรารอการคืนมัดจำมาแล้วสองสัปดาห์",
            "answer_en": "had been waiting",
            "explain_en": "Past Perfect Continuous (duration to a past point).",
            "explain_th": "คำอธิบาย: PPC — ระยะเวลาก่อนถึงวันศุกร์."},
        {
            "en": "Tomorrow at 10, the tenants [[(will be signing)]] the renewal, and the agent [[(will be witnessing)]] it.",
            "th": "พรุ่งนี้สิบโมง ผู้เช่าจะกำลังเซ็นต่อสัญญา และเอเจนต์จะกำลังกระทำการเป็นพยาน",
            "answer_en": "will be signing + will be witnessing",
            "explain_en": "two Future Continuous actions.",
            "explain_th": "คำอธิบาย: Future Continuous สองเหตุการณ์พร้อมกันในอนาคต."},
        {
            "en": "Since moving in, we [[(have dealt)]] with three repair requests.",
            "th": "ตั้งแต่ย้ายเข้า เราได้จัดการคำขอซ่อมสามครั้งแล้ว",
            "answer_en": "have dealt",
            "explain_en": "Present Perfect (experience up to now).",
            "explain_th": "คำอธิบาย: Present Perfect — ประสบการณ์ถึงปัจจุบัน."},
        {
            "en": "At the viewing, the landlord [[(was explaining)]] the house rules while we [[(were taking)]] notes.",
            "th": "ตอนเข้าชม เจ้าของบ้านกำลังอธิบายกฎของบ้าน ในขณะที่พวกเรากำลังจดบันทึก",
            "answer_en": "was explaining + were taking",
            "explain_en": "two Past Continuous actions.",
            "explain_th": "คำอธิบาย: Past Continuous สองเหตุการณ์พร้อมกัน."},
        {
            "en": "By the end of this month, we [[(will have paid)]] rent for twelve months in total.",
            "th": "ภายในสิ้นเดือนนี้ เราจะจ่ายค่าเช่าครบสิบสองเดือนแล้ว",
            "answer_en": "will have paid", "explain_en": "Future Perfect.",
            "explain_th": "คำอธิบาย: Future Perfect — เสร็จสิ้นก่อนสิ้นเดือน."},
        {
            "en": "We [[(have been discussing)]] a rent discount because the elevator [[(has been breaking)]] down lately.",
            "th": "เรากำลังถกเรื่องส่วนลดค่าเช่า เพราะลิฟต์ช่วงนี้เสียบ่อย",
            "answer_en": "have been discussing + has been breaking",
            "explain_en": "PPC + PPC.",
            "explain_th": "คำอธิบาย: PPC ทั้งสอง — ต่อเนื่องถึงปัจจุบัน."},
        {
            "en": "Last year we [[(were subletting)]] the spare room while we [[(were traveling)]] abroad.",
            "th": "ปีที่แล้ว เราปล่อยเช่าห้องว่างขณะที่เรากำลังเดินทางต่างประเทศ",
            "answer_en": "were subletting + were traveling",
            "explain_en": "two Past Continuous actions.",
            "explain_th": "คำอธิบาย: Past Continuous สองเหตุการณ์คู่ขนาน."},
        {
            "en": "If the roofers finish by noon, maintenance [[(will have scheduled)]] the inspection for the afternoon.",
            "th": "ถ้าช่างมุงหลังคาทำเสร็จภายในเที่ยง ทีมซ่อมบำรุงจะได้กำหนดการตรวจไว้สำหรับตอนบ่ายแล้ว",
            "answer_en": "will have scheduled",
            "explain_en": "Future Perfect after condition (by noon).",
            "explain_th": "คำอธิบาย: Future Perfect — เสร็จภายในเวลาที่กำหนดหลังเงื่อนไข."},
        {
            "en": "They [[(had been living)]] with a roommate before they [[(moved)]] into a studio.",
            "th": "ก่อนย้ายไปอยู่สตูดิโอ พวกเขาอยู่กับรูมเมทมาก่อน",
            "answer_en": "had been living + moved",
            "explain_en": "PPC + Past Simple.",
            "explain_th": "คำอธิบาย: PPC (ระยะเวลา) ก่อน Past Simple (เหตุการณ์ย้าย)."},
        {
            "en": "By the time the notice period ends, we [[(will have arranged)]] movers and [[(will have cleaned)]] the flat.",
            "th": "เมื่อครบกำหนดแจ้งออก เราจะจัดการเรื่องคนขนของและทำความสะอาดห้องเสร็จแล้ว",
            "answer_en": "will have arranged + will have cleaned",
            "explain_en": "two Future Perfect actions.",
            "explain_th": "คำอธิบาย: Future Perfect สองเหตุการณ์เสร็จก่อนกำหนด."},
        {
            "en": "For months the landlord [[(had been ignoring)]] our emails, but finally he [[(replied)]] last week.",
            "th": "หลายเดือนเจ้าของบ้านไม่สนอีเมลของเรา แต่ในที่สุดเขาตอบเมื่อสัปดาห์ก่อน",
            "answer_en": "had been ignoring + replied",
            "explain_en": "PPC + Past Simple.",
            "explain_th": "คำอธิบาย: PPC (ต่อเนื่องในอดีต) แล้วเกิด Past Simple (เหตุการณ์ครั้งเดียว)."},
        {
            "en": "At 7 p.m. tomorrow, I [[(will be meeting)]] a potential roommate who [[(has applied)]] for the room.",
            "th": "พรุ่งนี้หนึ่งทุ่ม ฉันจะกำลังพบรูมเมทที่เป็นไปได้ซึ่งได้สมัครไว้",
            "answer_en": "will be meeting + has applied",
            "explain_en": "Future Continuous + Present Perfect.",
            "explain_th": "คำอธิบาย: นัดพบในอนาคต + ผลที่เกิดขึ้นแล้ว."},
        {
            "en": "Since last winter, the building [[(has been undergoing)]] major renovations.",
            "th": "ตั้งแต่ฤดูหนาวที่แล้ว อาคารอยู่ระหว่างการปรับปรุงใหญ่",
            "answer_en": "has been undergoing",
            "explain_en": "Present Perfect Continuous.",
            "explain_th": "คำอธิบาย: PPC — ต่อเนื่องตั้งแต่ฤดูหนาวที่แล้ว."},
        {
            "en": "By last Tuesday, we [[(had submitted)]] all documents and [[(had scheduled)]] the handover.",
            "th": "ถึงวันอังคารที่แล้ว เราส่งเอกสารทั้งหมดและนัดส่งมอบเรียบร้อยแล้ว",
            "answer_en": "had submitted + had scheduled",
            "explain_en": "two Past Perfect actions.",
            "explain_th": "คำอธิบาย: Past Perfect สองเหตุการณ์สำเร็จก่อนวันอังคาร."},
        {
            "en": "Right now, the tenants [[(are discussing)]] an extension while the agent [[(is drafting)]] a new clause.",
            "th": "ตอนนี้ผู้เช่ากำลังหารือเรื่องต่อสัญญา ขณะที่เอเจนต์กำลังกdraft ข้อใหม่",
            "answer_en": "are discussing + is drafting",
            "explain_en": "two Present Continuous actions.",
            "explain_th": "คำอธิบาย: Present Continuous สองเหตุการณ์ขณะนี้."},
        {
            "en": "Every year, we [[(review)]] the lease terms and [[(negotiate)]] the rent.",
            "th": "ทุกปี เราทบทวนข้อกำหนดสัญญาและต่อรองค่าเช่า",
            "answer_en": "review + negotiate",
            "explain_en": "Present Simple (repeated actions).",
            "explain_th": "คำอธิบาย: Present Simple — กิจกรรมที่ทำเป็นประจำ."},
        {
            "en": "By next quarter, the property manager [[(will have implemented)]] new house rules.",
            "th": "ภายในไตรมาสหน้า ผู้จัดการทรัพย์สินจะได้บังคับใช้กฎใหม่แล้ว",
            "answer_en": "will have implemented",
            "explain_en": "Future Perfect.",
            "explain_th": "คำอธิบาย: Future Perfect — เสร็จภายในไตรมาสหน้า."},
        {
            "en": "When the plumber arrived, we [[(had already shut off)]] the water and [[(were waiting)]] in the lobby.",
            "th": "เมื่อช่างประปามาถึง เราได้ปิดน้ำไว้แล้วและกำลังรออยู่ในโถง",
            "answer_en": "had already shut off + were waiting",
            "explain_en": "Past Perfect + Past Continuous.",
            "explain_th": "คำอธิบาย: Past Perfect (เสร็จแล้ว) + Past Continuous (กำลังรอขณะนั้น)."},
        {
            "en": "By next week, I [[(will have been comparing)]] rental listings for a month.",
            "th": "ถึงสัปดาห์หน้า ฉันจะเปรียบเทียบประกาศเช่ามาเป็นเวลา 1 เดือนแล้ว",
            "answer_en": "will have been comparing",
            "explain_en": "Future Perfect Continuous.",
            "explain_th": "คำอธิบาย: FPC — ระยะเวลาถึงสัปดาห์หน้า."},
    ],
    # Vocab-блок (слово — RU — TH), плюс упражнения по вокабу (с ответами)
    "vocab": [
        {"en": "tenant", "ru": "арендатор", "th": "ผู้เช่า", "emoji": "👤"},
        {"en": "landlord / landlady", "ru": "арендодатель / хозяйка",
         "th": "เจ้าของบ้าน", "emoji": "🧑‍💼"},
        {"en": "lease / rental agreement", "ru": "договор аренды",
         "th": "สัญญาเช่า", "emoji": "📄"},
        {"en": "security deposit", "ru": "страховой депозит", "th": "เงินมัดจำ",
         "emoji": "💰"},
        {"en": "utilities (water, gas, electricity)",
         "ru": "коммунальные услуги (вода, газ, электричество)",
         "th": "ค่าสาธารณูปโภค (น้ำ แก๊ส ไฟฟ้า)", "emoji": "💡"},
        {"en": "maintenance", "ru": "техническое обслуживание",
         "th": "การบำรุงรักษา", "emoji": "🛠️"},
        {"en": "repair request", "ru": "заявка на ремонт", "th": "คำขอซ่อม",
         "emoji": "📝"},
        {"en": "property manager", "ru": "управляющий недвижимостью",
         "th": "ผู้จัดการทรัพย์สิน", "emoji": "🏢"},
        {"en": "real estate agent / realtor", "ru": "риелтор",
         "th": "ตัวแทนอสังหาริมทรัพย์", "emoji": "🏠"},
        {"en": "notice to vacate", "ru": "уведомление о выезде",
         "th": "หนังสือแจ้งย้ายออก", "emoji": "📬"},
        {"en": "move-in / move-out", "ru": "въезд / выезд",
         "th": "ย้ายเข้า / ย้ายออก", "emoji": "🚚"},
        {"en": "condition report / checklist", "ru": "акт приёма-передачи",
         "th": "แบบฟอร์มตรวจรับ / เช็กลิสต์", "emoji": "📋"},
        {"en": "inspection", "ru": "осмотр/проверка", "th": "การตรวจสอบ",
         "emoji": "🔎"},
        {"en": "keys / key handover", "ru": "ключи / передача ключей",
         "th": "กุญแจ / การส่งมอบกุญแจ", "emoji": "🔑"},
        {"en": "mailbox", "ru": "почтовый ящик", "th": "ตู้จดหมาย",
         "emoji": "📮"},
        {"en": "lock / change the lock", "ru": "замок / сменить замок",
         "th": "ล็อก / เปลี่ยนล็อก", "emoji": "🔒"},
        {"en": "furnished / unfurnished", "ru": "меблированный / без мебели",
         "th": "มีเฟอร์นิเจอร์ / ไม่มีเฟอร์นิเจอร์", "emoji": "🛋️"},
        {"en": "appliances", "ru": "бытовая техника", "th": "เครื่องใช้ไฟฟ้า",
         "emoji": "🔌"},
        {"en": "pet policy", "ru": "правила по животным",
         "th": "นโยบายสัตว์เลี้ยง", "emoji": "🐾"},
        {"en": "house rules", "ru": "правила дома", "th": "กฎของบ้าน",
         "emoji": "📘"},
        {"en": "sublet / sublease", "ru": "сдавать в субаренду",
         "th": "ให้เช่าช่วง", "emoji": "🔄"},
        {"en": "roommate / flatmate", "ru": "сосед по квартире",
         "th": "เพื่อนร่วมห้อง", "emoji": "👥"},
        {"en": "listing", "ru": "объявление (листинг)",
         "th": "ประกาศให้เช่า/ขาย", "emoji": "📢"},
        {"en": "viewing", "ru": "просмотр квартиры", "th": "การนัดชม",
         "emoji": "👀"},
        {"en": "addendum", "ru": "дополнительное соглашение",
         "th": "ภาคผนวกสัญญา", "emoji": "➕"},
        {"en": "notice period", "ru": "срок уведомления",
         "th": "ระยะเวลาแจ้งล่วงหน้า", "emoji": "⏳"},
        {"en": "rent increase / raise", "ru": "повышение аренды",
         "th": "การขึ้นค่าเช่า", "emoji": "📈"},
        {"en": "discount / rent reduction", "ru": "скидка / снижение аренды",
         "th": "ส่วนลด / การลดค่าเช่า", "emoji": "📉"},
        {"en": "handyman / contractor", "ru": "мастер / подрядчик",
         "th": "ช่างซ่อม / ผู้รับเหมา", "emoji": "👷"},
        {"en": "plumbing leak", "ru": "протечка сантехники",
         "th": "น้ำรั่ว (ระบบประปา)", "emoji": "💧"},
        {"en": "heating system", "ru": "отопительная система",
         "th": "ระบบทำความร้อน", "emoji": "♨️"},
        {"en": "smoke detector", "ru": "датчик дыма",
         "th": "เครื่องตรวจจับควัน", "emoji": "🚨"},
        {"en": "elevator / lift", "ru": "лифт", "th": "ลิฟต์", "emoji": "🛗"},
        {"en": "renovation / refurbishment", "ru": "ремонт/обновление",
         "th": "การปรับปรุง / บูรณะ", "emoji": "🏗️"},
        {"en": "refund (deposit refund)", "ru": "возврат (депозита)",
         "th": "การคืนเงิน (คืนเงินมัดจำ)", "emoji": "💸"},
        {"en": "movers / moving company",
         "ru": "грузчики / мувинговая компания", "th": "คนขนของ / บริษัทขนย้าย",
         "emoji": "🚚"},
        {"en": "add utilities to your name",
         "ru": "оформить коммуналку на себя", "th": "โอนสาธารณูปโภคเป็นชื่อคุณ",
         "emoji": "🧾"},
        {"en": "late fee", "ru": "штраф за просрочку", "th": "ค่าปรับล่าช้า",
         "emoji": "⏰"},
        {"en": "eviction", "ru": "выселение", "th": "การขับไล่ผู้เช่า",
         "emoji": "🚫"},
        {"en": "rent receipt", "ru": "квитанция об оплате аренды",
         "th": "ใบเสร็จค่าเช่า", "emoji": "🧾"},
        {"en": "monthly rent / due date",
         "ru": "месячная арендная плата / срок оплаты",
         "th": "ค่าเช่ารายเดือน / กำหนดชำระ", "emoji": "📆"},
        {"en": "garbage disposal / trash pickup", "ru": "вывоз мусора",
         "th": "การจัดการขยะ / การเก็บขยะ", "emoji": "🗑️"},
        {"en": "parking spot / permit", "ru": "парковочное место / разрешение",
         "th": "ที่จอดรถ / ใบอนุญาต", "emoji": "🅿️"},
        {"en": "common area", "ru": "общая зона", "th": "พื้นที่ส่วนกลาง",
         "emoji": "🏘️"},
        {"en": "noise complaint", "ru": "жалоба на шум",
         "th": "การร้องเรียนเรื่องเสียงดัง", "emoji": "🔊"},
    ],
    "vocab_ex": [
        {"en": "Use [[tenant]] in a sentence about on-time payment.",
         "ru": "Используй «tenant» в предложении про оплату вовремя.",
         "th": "ใช้ tenant ในประโยคเกี่ยวกับการจ่ายตรงเวลา",
         "answer_en": "The tenant has always paid on time.",
         "explain_en": "Present Perfect for life pattern."},
        {"en": "Make a question with [[landlord]] about repairs next week.",
         "ru": "Сделай вопрос к арендодателю о ремонте на след. неделе.",
         "th": "ตั้งคำถามถึงเจ้าของบ้านเรื่องซ่อมสัปดาห์หน้า",
         "answer_en": "Will the landlord repair the sink next week?",
         "explain_en": "Future Simple question."},
        {"en": "Use [[lease]] to state a rule (present).",
         "ru": "Используй «lease», чтобы обозначить правило (настоящее).",
         "th": "ใช้ lease เพื่อระบุข้อกำหนด (ปัจจุบัน)",
         "answer_en": "The lease requires quiet hours after 10 pm.",
         "explain_en": "Present Simple rule."},
        {"en": "Use [[security deposit]] (completed in the past).",
         "ru": "Используй «security deposit» (завершено в прошлом).",
         "th": "ใช้ security deposit (เกิดขึ้นในอดีต)",
         "answer_en": "They paid the security deposit yesterday.",
         "explain_en": "Past Simple."},
        {"en": "Write a sentence with [[utility bills]] (result by now).",
         "ru": "Предложение с «utility bills» (результат к настоящему).",
         "th": "ประโยคกับ utility bills (ผลถึงปัจจุบัน)",
         "answer_en": "We have paid all utility bills.",
         "explain_en": "Present Perfect."},
        {"en": "Use [[maintenance]] (duration up to now).",
         "ru": "Используй «maintenance» (длится до настоящего).",
         "th": "ใช้ maintenance (ต่อเนื่องถึงปัจจุบัน)",
         "answer_en": "The building has been under maintenance for a week.",
         "explain_en": "Present Perfect Continuous."},
        {"en": "Use [[inspection]] with a specific time in the past.",
         "ru": "«inspection» со временем в прошлом.",
         "th": "inspection กับเวลาที่ชัดเจนในอดีต",
         "answer_en": "The inspection started at 9 am yesterday.",
         "explain_en": "Past Simple."},
        {"en": "Ask about [[notice]] period politely.",
         "ru": "Спроси вежливо о сроке «notice».",
         "th": "ถามอย่างสุภาพเกี่ยวกับระยะเวลา notice",
         "answer_en": "Could you tell me the notice period, please?",
         "explain_en": "Polite question."},
        {"en": "Use [[move-in]] with a future plan in progress.",
         "ru": "«move-in» с будущим процессом.",
         "th": "move-in กับเหตุการณ์กำลังเกิดในอนาคต",
         "answer_en": "We will be moving in next Monday afternoon.",
         "explain_en": "Future Continuous."},
        {"en": "Use [[move-out]] with earlier-past meaning.",
         "ru": "«move-out» с предпрошедшим значением.",
         "th": "move-out กับเหตุการณ์ก่อนอดีต",
         "answer_en": "We had moved out before the renovation started.",
         "explain_en": "Past Perfect."},
        {"en": "Use [[furnishings]] with a present rule.",
         "ru": "«furnishings» с правилом в настоящем.",
         "th": "furnishings กับกฎในปัจจุบัน",
         "answer_en": "Furnishings must remain in the unit.",
         "explain_en": "Present Simple obligation."},
        {"en": "Use [[inventory list]] with completion by a future time.",
         "ru": "«inventory list» — завершённость к будущему времени.",
         "th": "inventory list เสร็จก่อนเวลาอนาคต",
         "answer_en": "We will have completed the inventory list by 5 pm.",
         "explain_en": "Future Perfect."},
        {"en": "Use [[keys]] in a passive past sentence.",
         "ru": "«keys» в пассивном в прошедшем.",
         "th": "keys ในประโยคถูกกระทำ (อดีต)",
         "answer_en": "The keys were handed over yesterday.",
         "explain_en": "Past Simple passive."},
        {"en": "Use [[parking permit]] in a present rule sentence.",
         "ru": "«parking permit» в предложении-правиле.",
         "th": "parking permit ในประโยคกฎ",
         "answer_en": "A parking permit is required for overnight parking.",
         "explain_en": "Present Simple rule."},
        {"en": "Use [[house rules]] with a polite request.",
         "ru": "«house rules» с вежливой просьбой.",
         "th": "house rules กับคำขอสุภาพ",
         "answer_en": "Please follow the house rules during quiet hours.",
         "explain_en": "Polite imperative."},
        {"en": "Use [[renewal]] with a decision now.",
         "ru": "«renewal» — решение сейчас.",
         "th": "renewal กับการตัดสินใจตอนนี้",
         "answer_en": "I will renew the lease this week.",
         "explain_en": "Future Simple decision."},
        {"en": "Use [[damage]] with duration before a past point.",
         "ru": "«damage» — длительность до прошлого момента.",
         "th": "damage ระยะเวลาก่อนอดีต",
         "answer_en": "The ceiling had been damaged for months before repairs.",
         "explain_en": "Past Perfect Continuous."},
        {"en": "Use [[repairs]] with duration until a future point.",
         "ru": "«repairs» — длительность к будущему моменту.",
         "th": "repairs ระยะเวลาถึงอนาคต",
         "answer_en": "By Friday, repairs will have been ongoing for two weeks.",
         "explain_en": "Future Perfect Continuous."},
    ],
    "exit_hw": [
        {"en": "Make 3 sentences using [[termination]] (any tenses).",
         "ru": "Сделай 3 предложения с «termination» (любой tense).",
         "th": "เขียน 3 ประโยคโดยใช้ termination (กาลใดก็ได้)",
         "answer_en": "e.g., We terminated the lease; It has been terminated; We will terminate it next month.",
         "explain_en": "Any correct tense/context is acceptable."},
        {"en": "Write a polite email line about [[late fee]] (present rule).",
         "ru": "Напиши вежливую строку про «late fee» (правило).",
         "th": "เขียนประโยคสุภาพเกี่ยวกับ late fee (กฎปัจจุบัน)",
         "answer_en": "Please note that a late fee applies after the 5th of each month.",
         "explain_en": "Present Simple rule."}
    ]
}

# ---------- Сборка DOCX (answers version) ----------
doc = new_doc()
add_title(doc, content["title"] + " — Answers")

# Explanation
add_section_title(doc, "👩‍🏫", "Explanation")
for b in content["explanation"]:
    p = doc.add_paragraph()
    r = p.add_run(b["title"])
    r.font.bold = True
    r.font.color.rgb = RGBColor(180, 90, 0)  # тёмно-оранжевый, читаемый
    examples_block(doc, b.get("examples", []), b.get("i", "1"))
    if b.get("ru"): line_ru(doc, b["ru"])
    if b.get("th"): line_th(doc, b["th"])
    add_blank(doc)

# Practice — Answers
doc.add_page_break()
add_section_title(doc, "🧠", "Practice — Answers")
for i, ex in enumerate(content["practice"], 1):
    add_exercise(doc, i, ex["en"], ex.get("ru"), ex.get("th"))
    add_answer_block(doc, ex["answer_en"], ex["explain_en"],
                     ex.get("explain_th"))

# Vocabulary (word bank)
doc.add_page_break()
add_section_title(doc, content.get("theme_emoji", "🧰"),
                  content.get("vocab_title", "Vocabulary"))
letters = "abcdefghijklmnopqrstuvwxyz"
for i, w in enumerate(content["vocab"]):
    letter = letters[i] if i < 26 else letters[i - 26] * 2
    add_word_bank_item(doc, letter, w["en"], w.get("ru"), w.get("th"),
                       w.get("emoji"))

# Vocabulary Exercises — Answers
doc.add_page_break()
add_section_title(doc, "🛄", "Vocabulary Exercises — Answers")
for i, ex in enumerate(content["vocab_ex"], 1):
    add_exercise(doc, i, ex["en"], ex.get("ru"), ex.get("th"))
    add_answer_block(doc, ex["answer_en"], ex["explain_en"])

# Exit — Answers
doc.add_page_break()
add_section_title(doc, "🧾", "Exit check & Homework — Answers")
for i, ex in enumerate(content["exit_hw"], 1):
    add_exercise(doc, i, ex["en"], ex.get("ru"), ex.get("th"))
    add_answer_block(doc, ex["answer_en"], ex["explain_en"])

out_path = "cha_test_12tenses_rentals_with_answers_v4.docx"
doc.save(out_path)
print(f"OK -> {out_path}")
