import argparse
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

# ---------- Цвета и шрифты ----------
BLACK = RGBColor(0, 0, 0)
DARK_RED = RGBColor(139, 0, 0)
DARK_GREEN = RGBColor(0, 100, 0)
PURPLE = RGBColor(102, 0, 153)
THAI_FONT_NAME = "Noto Sans Thai"

OUT_NAME = "cha_lesson_3_w_questions_v1.docx"
SRC_NAME = "cha_lesson_3_w-questions_v8.docx"

BLOCK_TITLES = {
    "🎓 Lesson 3 — W-Questions — Vocabulary: Student Graduation",
    "👩‍🏫 Explanation",
    "✍️ Examples:",
    "🧠 Practice",
    "🎓 Vocabulary (Student Graduation)",
    "🧺 Word bank:",
    "🎓 Vocabulary Exercises",
    "🧾 Exit check & Homework",
    "🧾 Exit check (5 quick items):",
}


def new_doc():
    doc = Document()
    for s in doc.sections:
        s.page_height = Cm(29.7)
        s.page_width = Cm(21.0)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        fp = s.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = fp.add_run("© Cha 2025 · Page ")
        run1.font.size = Pt(9)
        run1.font.color.rgb = BLACK
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run2 = fp.add_run()
        run2._r.append(fld)
    return doc


def clone_run(dst_p, src_run):
    r = dst_p.add_run(src_run.text)
    try:
        r.font.bold = src_run.font.bold
        r.font.italic = src_run.font.italic
        r.font.underline = src_run.font.underline
        r.font.size = src_run.font.size
        r.font.all_caps = src_run.font.all_caps
        if src_run.font.color and src_run.font.color.rgb:
            r.font.color.rgb = src_run.font.color.rgb
    except Exception:
        pass
    return r


def clone_paragraph(dst_doc: Document, src_p):
    p = dst_doc.add_paragraph()
    p.alignment = src_p.alignment
    for run in src_p.runs:
        clone_run(p, run)
    return p


def line_ru(doc: Document, txt: str, size=11):
    p = doc.add_paragraph()
    r = p.add_run(f"({txt})")
    r.font.italic = True
    r.font.color.rgb = DARK_RED
    r.font.size = Pt(size)


def line_th(doc: Document, txt: str, size=11):
    p = doc.add_paragraph()
    r = p.add_run(f"({txt})")
    r.font.italic = True
    r.font.color.rgb = DARK_GREEN
    r.font.size = Pt(size)
    r.font.name = THAI_FONT_NAME


def collect_highlight_tokens(src_p) -> list:
    tokens = []
    for run in src_p.runs:
        t = run.text or ""
        for m in re.finditer(r"[A-Z][A-Z ]+[A-Z]", t):
            tok = m.group(0).strip()
            if tok not in tokens:
                tokens.append(tok)
        try:
            if run.font and run.font.underline and not any(
                    ch.isupper() for ch in t):
                w = (t or "").strip()
                if 0 < len(w) <= 15 and w not in tokens:
                    tokens.append(w)
        except Exception:
            pass
    tokens.sort(key=len, reverse=True)
    return tokens


def add_ru_mapped_line_with_highlights(doc: Document, src_p, ru_text: str):
    p = doc.add_paragraph()
    r0 = p.add_run("(")
    r0.font.italic = True
    r0.font.color.rgb = DARK_RED
    hi = collect_highlight_tokens(src_p)
    s = ru_text or ""
    i = 0
    while i < len(s):
        hit_pos = None
        hit_tok = None
        for tok in hi:
            j = s.find(tok, i)
            if j != -1 and (hit_pos is None or j < hit_pos):
                hit_pos = j
                hit_tok = tok
        if hit_pos is None:
            r = p.add_run(s[i:])
            r.font.italic = True
            r.font.color.rgb = DARK_RED
            break
        if hit_pos > i:
            r = p.add_run(s[i:hit_pos])
            r.font.italic = True
            r.font.color.rgb = DARK_RED
        r2 = p.add_run(s[hit_pos:hit_pos + len(hit_tok)])
        r2.font.color.rgb = BLACK
        r2.font.bold = True
        r2.font.underline = True
        r2.font.italic = False
        i = hit_pos + len(hit_tok)
    rz = p.add_run(")")
    rz.font.italic = True
    rz.font.color.rgb = DARK_RED


def add_th_mapped_line_with_highlights(doc: Document, src_p, th_text: str):
    p = doc.add_paragraph()
    r0 = p.add_run("(")
    r0.font.italic = True
    r0.font.color.rgb = DARK_GREEN
    hi = collect_highlight_tokens(src_p)
    s = th_text or ""
    i = 0
    while i < len(s):
        hit_pos = None
        hit_tok = None
        for tok in hi:
            j = s.find(tok, i)
            if j != -1 and (hit_pos is None or j < hit_pos):
                hit_pos = j
                hit_tok = tok
        if hit_pos is None:
            r = p.add_run(s[i:])
            r.font.italic = True
            r.font.color.rgb = DARK_GREEN
            r.font.name = THAI_FONT_NAME
            break
        if hit_pos > i:
            r = p.add_run(s[i:hit_pos])
            r.font.italic = True
            r.font.color.rgb = DARK_GREEN
            r.font.name = THAI_FONT_NAME
        r2 = p.add_run(s[hit_pos:hit_pos + len(hit_tok)])
        r2.font.color.rgb = BLACK
        r2.font.bold = True
        r2.font.underline = True
        r2.font.italic = False
        r2.font.name = THAI_FONT_NAME
        i = hit_pos + len(hit_tok)
    rz = p.add_run(")")
    rz.font.italic = True
    rz.font.color.rgb = DARK_GREEN


def norm_exact(s: str) -> str:
    s = (s or "").strip()
    s = s.replace("\u2011", "-")
    s = s.replace("\u2013", "-")
    s = s.replace("\u2014", "-")
    s = s.replace("\u2212", "-")
    s = re.sub(r"^\s*\d+(?:\.\d+)*[\)\.]?\s+", "", s)
    s = re.sub(r"^[\u2022\-\u2013\u2014]\s+", "", s)
    return re.sub(r"\s+", " ", s)


def strip_list_markers(s: str) -> str:
    return re.sub(r"^[\u2022\-\u2013\u2014]\s+", "", (s or "").strip())


def clean_vocab_en_term(s: str) -> str:
    """Очищает EN-термин Word bank: убирает литерную нумерацию (a.), эмодзи; оставляет латиницу/пробелы/дефис/скобки."""
    s = (s or "").strip()
    s = s.replace("\u2011", "-").replace("\u2013", "-").replace("\u2014",
                                                                "-").replace(
        "\u2212", "-")
    s = re.sub(r"^[A-Za-z]\.[\s]+", "", s)
    s = re.sub(r"[^A-Za-z()\-\s]", "", s)
    s = re.sub(r"\s+", " ", s).strip().lower()
    return s


def load_translations_from_source(path: str) -> dict:
    """Парсит файл перевода: EN строка + (RU) + (TH) как отдельные строки. Пропускает блок Word bank."""
    if not path or not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        lines = [ln.rstrip("\n") for ln in f]
    tr = {}
    section = None
    # защита от дублей: по индексу исходного параграфа
    added_ru_idx = set()
    added_th_idx = set()
    i = 0
    while i < len(lines):
        L = lines[i].strip()
        if not L:
            i += 1
            continue
        low = L.lower()
        if L in BLOCK_TITLES:
            if "vocabulary (student graduation)" in low:
                section = "vocab"
            elif "vocabulary exercises" in low:
                section = "vocab_ex"
            elif "practice" in low:
                section = "practice"
            elif "exit check" in low or "homework" in low:
                section = "exit"
            elif "explanation" in low:
                section = "expl"
            i += 1
            continue
        # пропускаем word bank — там перевод в одну строку
        if section == "vocab":
            i += 1
            continue
        if not L.startswith("("):
            # Собираем до нескольких следующих строк в скобках и классифицируем язык по алфавиту
            ru = th = None
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("("):
                raw = lines[j].strip()
                val = raw[1:-1] if (
                        raw.startswith("(") and raw.endswith(")")) else raw
                # Классификация по символам (тайский / кириллица)
                if re.search(r"[\u0E00-\u0E7F]", val):  # Thai block
                    th = val  # берём последнее встреченное TH
                elif re.search(r"[\u0400-\u04FF]", val):  # Cyrillic
                    ru = val  # берём последнее встреченное RU
                else:
                    # если не удалось классифицировать — не учитываем
                    pass
                j += 1
            if ru or th:
                base_key = norm_exact(L)
                pair = {"ru": ru, "th": th}
                tr[base_key] = pair
                alt = norm_exact(strip_list_markers(L))
                if alt != base_key:
                    tr[alt] = pair
                i = j
                continue
        i += 1
    return tr


def load_wordbank_from_source(path: str) -> dict:
    """Парсит блок Word bank из текстового источника. Возвращает dict по ключу EN-термина -> {ru, th}."""
    if not path or not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        lines = [ln.rstrip("\n") for ln in f]
    wb = {}
    section = None
    for L in lines:
        S = L.strip()
        if not S:
            continue
        low = S.lower()
        if S in BLOCK_TITLES:
            if "vocabulary (student graduation)" in low:
                section = "vocab"
            else:
                section = None
            continue
        if section != "vocab":
            continue
        # ожидаем формат: a. <emoji?> EN — RU — TH
        if re.match(r"^[A-Za-z]\.", S) and " — " in S:
            parts = re.split(r"\s—\s", S, maxsplit=2)
            if len(parts) >= 2:
                left = parts[0]
                ru = parts[1] if len(parts) >= 2 else None
                th = parts[2] if len(parts) >= 3 else None
                # ключи
                keys = set()
                keys.add(norm_exact(left))
                keys.add(norm_exact(re.sub(r"^[A-Za-z]\.[\s]+", "", left)))
                cv = clean_vocab_en_term(left)
                if cv:
                    keys.add(cv)
                for k in keys:
                    wb[k] = {"ru": ru, "th": th}
    return wb


def load_answers_from_source(path: str) -> dict:
    if not path or not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        lines = [ln.rstrip("\n") for ln in f]
    i = 0
    ans = {}
    while i < len(lines):
        en = lines[i].strip()
        if not en:
            i += 1
            continue
        if en in BLOCK_TITLES:
            i += 1
            continue
        if i + 2 <= len(lines) - 1 and lines[i + 1].lstrip().startswith(
                "Answer:") and lines[i + 2].lstrip().startswith("คำตอบ:"):
            key = norm_exact(en)
            a_en = lines[i + 1].strip()
            a_th = lines[i + 2].strip()
            ans.setdefault(key, []).extend([a_en, a_th])
            i += 3
            if i < len(lines) and not lines[i].strip():
                i += 1
            continue
        i += 1
    return ans


def build():
    parser = argparse.ArgumentParser()
    parser.add_argument("--with-ru", dest="with_ru", action="store_true",
                        default=True)
    parser.add_argument("--no-ru", dest="with_ru", action="store_false")
    parser.add_argument("--with-th", dest="with_th", action="store_true",
                        default=True)
    parser.add_argument("--no-th", dest="with_th", action="store_false")
    parser.add_argument("--translations-source", type=str,
                        default="lesson3_translations_source.txt")
    parser.add_argument("--with-answers", dest="with_answers",
                        action="store_true", default=False)
    parser.add_argument("--answers-source", type=str,
                        default="lesson3_answers_source.txt")
    args = parser.parse_args()

    # Проверки
    if not os.path.exists(SRC_NAME):
        raise FileNotFoundError(f"Source DOCX not found: {SRC_NAME}")
    if not args.translations_source or not os.path.exists(
            args.translations_source):
        raise FileNotFoundError(
            f"Translations source not found: {args.translations_source}")

    # Грузим маппинги
    tr_map = load_translations_from_source(args.translations_source)
    wb_map = load_wordbank_from_source(args.translations_source)
    ans_map = {}
    if args.with_answers and args.answers_source and os.path.exists(
            args.answers_source):
        ans_map = load_answers_from_source(args.answers_source)

    # База и выходной документ
    src = Document(SRC_NAME)
    out = new_doc()

    section = None
    total = len(src.paragraphs)
    for idx, p in enumerate(src.paragraphs, 1):
        text = p.text or ""

        t = text.strip().lower()
        if "vocabulary (student graduation)" in t:
            section = "vocab"
        elif "vocabulary exercises" in t:
            section = "vocab_ex"
        elif "practice" in t:
            section = "practice"
        elif "exit check" in t or "homework" in t:
            section = "exit"
        elif "explanation" in t or "examples" in t:
            section = "expl"

        stripped = text.strip()
        if not stripped:
            continue
        if stripped in BLOCK_TITLES:
            # Заголовки переносим как есть
            clone_paragraph(out, p)
            continue

        # Если текущий абзац уже является переводной строкой в скобках — не добавляем ничего
        if stripped.startswith("(") and stripped.endswith(")"):
            # Не переносим исходные переводные строки из базы — мы генерим свои
            continue

        # На этом этапе переносим сам EN-абзац в выход
        new_p = clone_paragraph(out, p)

        # Word bank: дописываем RU/TH в ту же строку
        if section == "vocab" and re.match(r"^[A-Za-z]\.[\s]+", stripped):
            # получить ключи поиска
            left = stripped.split(" — ", 1)[0]
            keys = [
                norm_exact(left),
                norm_exact(re.sub(r"^[A-Za-z]\.[\s]+", "", left)),
                clean_vocab_en_term(left),
                clean_vocab_en_term(re.sub(r"^[A-Za-z]\.[\s]+", "", left)),
            ]
            val = None
            for k in keys:
                if not k:
                    continue
                v = wb_map.get(k)
                if v:
                    val = v
                    break
            if val:
                ru = val.get("ru")
                th = val.get("th")
                if args.with_ru and ru:
                    rr_sep = new_p.add_run(" — ")
                    rr_sep.font.italic = True
                    rr_sep.font.color.rgb = DARK_RED
                    rr_run = new_p.add_run(ru)
                    rr_run.font.italic = True
                    rr_run.font.color.rgb = DARK_RED
                if args.with_th and th:
                    th_sep = new_p.add_run(" — ")
                    th_sep.font.italic = True
                    th_sep.font.color.rgb = DARK_GREEN
                    trun = new_p.add_run(th)
                    trun.font.italic = True
                    trun.font.color.rgb = DARK_GREEN
                    trun.font.name = THAI_FONT_NAME
            continue

        # Контентные строки: добавляем переводы
        key = norm_exact(text)

        # Для блока Exit check — особый формат: метки "— RU:" / "— TH:" вместо строк в скобках
        if section == "exit":
            if args.with_ru and idx not in added_ru_idx:
                ru_txt = tr_map.get(key, {}).get("ru")
                if ru_txt:
                    pr = out.add_paragraph()
                    rr = pr.add_run(f"— RU: {ru_txt}")
                    # стандартный стиль (чёрный, без курсива)
                    added_ru_idx.add(idx)
            if args.with_th and idx not in added_th_idx:
                th_txt = tr_map.get(key, {}).get("th")
                if th_txt:
                    pt = out.add_paragraph()
                    rt = pt.add_run(f"— TH: {th_txt}")
                    try:
                        rt.font.name = THAI_FONT_NAME
                    except Exception:
                        pass
                    added_th_idx.add(idx)
        else:
            # Остальные секции — как в уроке 4 (строки в скобках с зеркалированием)
            if args.with_ru and idx not in added_ru_idx:
                ru_txt = tr_map.get(key, {}).get("ru")
                if ru_txt:
                    add_ru_mapped_line_with_highlights(out, p, ru_txt)
                    added_ru_idx.add(idx)
            if args.with_th and idx not in added_th_idx:
                th_txt = tr_map.get(key, {}).get("th")
                if th_txt:
                    add_th_mapped_line_with_highlights(out, p, th_txt)
                    added_th_idx.add(idx)

        # Ответы — строго после переводов
        if args.with_answers and section in ("practice", "vocab_ex", "exit"):
            a = ans_map.get(key)
            if a:
                for line in a:
                    ap = out.add_paragraph()
                    ar = ap.add_run(line)
                    ar.font.color.rgb = PURPLE

    out.save(OUT_NAME)
    print("[lesson3] Saved:", OUT_NAME)


if __name__ == "__main__":
    build()
